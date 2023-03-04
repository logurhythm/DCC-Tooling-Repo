import docx

TEMPLATE_LIST = []
NEW_LIST = []
doc = docx.Document()

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    NEW_LIST.append(fullText)
    return '\n'.join(fullText)

def template_list_loader(CompanyName, JobTitle):
    for i in TEMPLATE_LIST:
        getText(i)
    for i in NEW_LIST:   
        doc.add_paragraph(i)
    doc.save(CompanyName + "-" + JobTitle + "_CL.docx")